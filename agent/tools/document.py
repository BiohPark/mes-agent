"""
문서 처리 도구
- Excel: openpyxl (읽기/쓰기/행 추가/셀 메모)
- Word: python-docx (읽기/내용 추가) + OpenXML (검토 메모/수정 추적)
- PDF: pdfplumber (텍스트 추출, 읽기 전용)
- PowerPoint: OpenXML (슬라이드 텍스트/발표자 노트)
- 텍스트 파일: 내장 (읽기/쓰기/추가)
"""

import json
import os
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path

_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
_A = 'http://schemas.openxmlformats.org/drawingml/2006/main'


# ── Excel ─────────────────────────────────────────────────────

def read_excel(path: str, sheet: int = 0, max_rows: int = 200) -> str:
    """Excel 파일을 읽어 JSON 배열로 반환합니다.
    sheet는 시트 번호(0부터) 또는 시트 이름입니다.
    첫 행이 헤더로 사용됩니다."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        if isinstance(sheet, int):
            ws = wb.worksheets[sheet]
        else:
            ws = wb[sheet]

        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            return json.dumps({"sheet": ws.title, "rows": [], "count": 0})

        headers = [str(c) if c is not None else f"col{i}" for i, c in enumerate(rows[0])]
        data = []
        for row in rows[1:max_rows + 1]:
            record = {headers[i]: (str(v) if v is not None else "") for i, v in enumerate(row)}
            data.append(record)
        return json.dumps({"sheet": ws.title, "rows": data, "count": len(data),
                           "headers": headers}, ensure_ascii=False)
    except Exception as e:
        return json.dumps({"error": str(e)})


def write_excel(path: str, data: list, sheet_name: str = "Sheet1") -> str:
    """데이터를 Excel 파일로 저장합니다.
    data는 딕셔너리 리스트이며, 키가 헤더가 됩니다."""
    try:
        import openpyxl
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = sheet_name

        if not data:
            wb.save(path)
            return json.dumps({"path": path, "rows": 0})

        headers = list(data[0].keys())
        ws.append(headers)
        for row in data:
            ws.append([row.get(h, "") for h in headers])

        Path(path).parent.mkdir(parents=True, exist_ok=True)
        wb.save(path)
        return json.dumps({"path": path, "rows": len(data), "headers": headers})
    except Exception as e:
        return json.dumps({"error": str(e)})


def append_excel_row(path: str, row_data: dict, sheet: int = 0) -> str:
    """Excel 파일의 마지막 행에 데이터를 추가합니다.
    파일이 없으면 새로 생성합니다."""
    try:
        import openpyxl
        if Path(path).exists():
            wb = openpyxl.load_workbook(path)
            ws = wb.worksheets[sheet] if isinstance(sheet, int) else wb[sheet]
            max_row = ws.max_row
            if max_row == 1:
                # 헤더만 있거나 빈 경우
                headers = [ws.cell(1, c).value for c in range(1, ws.max_column + 1)]
                ws.append([row_data.get(str(h), "") for h in headers])
            else:
                headers = [ws.cell(1, c).value for c in range(1, ws.max_column + 1)]
                ws.append([row_data.get(str(h), "") for h in headers])
        else:
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.append(list(row_data.keys()))
            ws.append(list(row_data.values()))

        Path(path).parent.mkdir(parents=True, exist_ok=True)
        wb.save(path)
        return json.dumps({"path": path, "added_row": row_data,
                           "total_rows": ws.max_row - 1})
    except Exception as e:
        return json.dumps({"error": str(e)})


def get_excel_sheet_names(path: str) -> str:
    """Excel 파일의 시트 목록을 반환합니다."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(path, read_only=True)
        return json.dumps({"sheets": wb.sheetnames})
    except Exception as e:
        return json.dumps({"error": str(e)})


# ── Word ──────────────────────────────────────────────────────

def read_word(path: str) -> str:
    """Word(.docx) 파일의 텍스트를 읽어 반환합니다."""
    try:
        import docx
        doc = docx.Document(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # 표 내용도 추출
        for table in doc.tables:
            for row in table.rows:
                row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_texts:
                    paragraphs.append(" | ".join(row_texts))
        return "\n".join(paragraphs) or "(내용 없음)"
    except Exception as e:
        return f"Word 읽기 실패: {e}"


def append_word(path: str, text: str, heading: str = "") -> str:
    """Word 파일에 내용을 추가합니다.
    heading이 있으면 제목 단락을 먼저 추가합니다. 파일이 없으면 새로 생성합니다."""
    try:
        import docx
        if Path(path).exists():
            doc = docx.Document(path)
        else:
            doc = docx.Document()
            Path(path).parent.mkdir(parents=True, exist_ok=True)

        if heading:
            doc.add_heading(heading, level=2)
        doc.add_paragraph(text)
        doc.save(path)
        return json.dumps({"path": path, "appended": text[:80]})
    except Exception as e:
        return json.dumps({"error": str(e)})


def _add_markdown_runs(paragraph, text: str) -> None:
    """단락에 **굵게**·`코드` 인라인 서식을 반영해 run을 추가한다."""
    import re
    # **bold** 와 `code` 를 토큰으로 분리
    tokens = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith('**') and tok.endswith('**'):
            run = paragraph.add_run(tok[2:-2])
            run.bold = True
        elif tok.startswith('`') and tok.endswith('`'):
            run = paragraph.add_run(tok[1:-1])
            run.font.name = 'Consolas'
        else:
            paragraph.add_run(tok)


def write_word(path: str, content: str, title: str = "") -> str:
    """마크다운 텍스트를 받아 서식이 살아있는 진짜 Word(.docx) 파일로 저장합니다.
    제목(#)·목록(-, 1.)·굵게(**)·표(|)를 Word 서식으로 변환합니다.
    경고: .docx 파일을 만들 때는 write_file 대신 반드시 이 도구를 사용해야 합니다.
    write_file로 .docx 경로에 텍스트를 쓰면 깨진 파일이 됩니다."""
    try:
        import docx
        doc = docx.Document()
        if title:
            doc.add_heading(title, level=0)

        lines = content.replace('\r\n', '\n').split('\n')
        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            # 빈 줄
            if not stripped:
                i += 1
                continue

            # 코드 펜스 ```
            if stripped.startswith('```'):
                i += 1
                code_lines = []
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                p = doc.add_paragraph()
                run = p.add_run('\n'.join(code_lines))
                run.font.name = 'Consolas'
                i += 1  # 닫는 ``` 건너뛰기
                continue

            # 표 (| a | b |) — 연속 행 수집
            if stripped.startswith('|') and stripped.endswith('|'):
                table_rows = []
                while i < len(lines) and lines[i].strip().startswith('|'):
                    cells = [c.strip() for c in lines[i].strip().strip('|').split('|')]
                    # 구분선(---) 행은 건너뜀
                    if not all(set(c) <= set('-: ') and c for c in cells):
                        table_rows.append(cells)
                    i += 1
                if table_rows:
                    cols = max(len(r) for r in table_rows)
                    table = doc.add_table(rows=0, cols=cols)
                    table.style = 'Light Grid Accent 1'
                    for ri, row in enumerate(table_rows):
                        wcells = table.add_row().cells
                        for ci in range(cols):
                            cell_text = row[ci] if ci < len(row) else ''
                            wp = wcells[ci].paragraphs[0]
                            _add_markdown_runs(wp, cell_text)
                            if ri == 0:
                                for r in wp.runs:
                                    r.bold = True
                continue

            # 제목 (#, ##, ###)
            if stripped.startswith('#'):
                level = len(stripped) - len(stripped.lstrip('#'))
                doc.add_heading(stripped.lstrip('#').strip(), level=min(level, 4))
                i += 1
                continue

            # 순서 없는 목록 (-, *)
            if stripped[:2] in ('- ', '* '):
                p = doc.add_paragraph(style='List Bullet')
                _add_markdown_runs(p, stripped[2:])
                i += 1
                continue

            # 순서 있는 목록 (1. 2. ...)
            import re as _re
            m = _re.match(r'^(\d+)\.\s+(.*)', stripped)
            if m:
                p = doc.add_paragraph(style='List Number')
                _add_markdown_runs(p, m.group(2))
                i += 1
                continue

            # 일반 단락
            p = doc.add_paragraph()
            _add_markdown_runs(p, stripped)
            i += 1

        Path(path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(path)
        return json.dumps({"path": path, "message": "Word 문서 저장 완료 (서식 변환됨)",
                           "chars": len(content)}, ensure_ascii=False)
    except Exception as e:
        return json.dumps({"error": str(e)})


# ── PDF ───────────────────────────────────────────────────────

def read_pdf(path: str, pages: str = "") -> str:
    """PDF 파일에서 텍스트를 추출합니다.
    pages 예시: '1', '1-3', '2,4,6' (1부터 시작). 생략하면 전체."""
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            total = len(pdf.pages)
            # 페이지 파싱
            if pages:
                page_nums = _parse_page_range(pages, total)
            else:
                page_nums = list(range(total))

            texts = []
            for i in page_nums:
                if 0 <= i < total:
                    page_text = pdf.pages[i].extract_text() or ""
                    texts.append(f"=== 페이지 {i+1} ===\n{page_text}")

            return "\n\n".join(texts) or "(추출된 텍스트 없음)"
    except Exception as e:
        return f"PDF 읽기 실패: {e}"


def _parse_page_range(pages: str, total: int) -> list:
    result = []
    for part in pages.replace(" ", "").split(","):
        if "-" in part:
            a, b = part.split("-", 1)
            result.extend(range(int(a) - 1, min(int(b), total)))
        else:
            result.append(int(part) - 1)
    return result


# ── Office 검토/메모 (OpenXML) ────────────────────────────────

def read_word_comments(path: str) -> str:
    """Word(.docx) 파일의 검토 메모(Comments)를 읽어 반환합니다.
    작성자·날짜·내용을 포함합니다."""
    try:
        with zipfile.ZipFile(path) as z:
            if 'word/comments.xml' not in z.namelist():
                return json.dumps({'comments': [], 'count': 0})
            ns = {'w': _W}
            root = ET.parse(z.open('word/comments.xml')).getroot()
            comments = []
            for c in root.findall('w:comment', ns):
                texts = [t.text for t in c.findall('.//w:t', ns) if t.text]
                comments.append({
                    'id': c.get(f'{{{_W}}}id'),
                    'author': c.get(f'{{{_W}}}author', ''),
                    'date': (c.get(f'{{{_W}}}date', '') or '')[:10],
                    'text': ''.join(texts)
                })
            return json.dumps({'comments': comments, 'count': len(comments)}, ensure_ascii=False)
    except Exception as e:
        return json.dumps({'error': str(e)})


def read_word_track_changes(path: str) -> str:
    """Word(.docx) 파일의 수정 추적(Track Changes) 삽입·삭제 내역을 반환합니다."""
    try:
        with zipfile.ZipFile(path) as z:
            if 'word/document.xml' not in z.namelist():
                return json.dumps({'error': '문서 파일 없음'})
            ns = {'w': _W}
            root = ET.parse(z.open('word/document.xml')).getroot()
            changes = []
            for ins in root.findall('.//w:ins', ns):
                texts = [t.text for t in ins.findall('.//w:t', ns) if t.text]
                changes.append({
                    'type': '삽입',
                    'author': ins.get(f'{{{_W}}}author', ''),
                    'date': (ins.get(f'{{{_W}}}date', '') or '')[:10],
                    'text': ''.join(texts)
                })
            for d in root.findall('.//w:del', ns):
                texts = [t.text for t in d.findall('.//w:delText', ns) if t.text]
                changes.append({
                    'type': '삭제',
                    'author': d.get(f'{{{_W}}}author', ''),
                    'date': (d.get(f'{{{_W}}}date', '') or '')[:10],
                    'text': ''.join(texts)
                })
            return json.dumps({'track_changes': changes, 'count': len(changes)}, ensure_ascii=False)
    except Exception as e:
        return json.dumps({'error': str(e)})


def read_excel_comments(path: str, sheet: str = '') -> str:
    """Excel(.xlsx) 파일의 셀 메모(Comments)를 읽어 반환합니다."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(path)
        sheets_to_check = [wb[sheet]] if sheet and sheet in wb.sheetnames else wb.worksheets
        results = {}
        for ws in sheets_to_check:
            comments = []
            for row in ws.iter_rows():
                for cell in row:
                    if cell.comment:
                        comments.append({
                            'cell': cell.coordinate,
                            'author': cell.comment.author or '',
                            'text': cell.comment.text or ''
                        })
            if comments:
                results[ws.title] = comments
        total = sum(len(v) for v in results.values())
        return json.dumps({'sheets': results, 'total_comments': total}, ensure_ascii=False)
    except Exception as e:
        return json.dumps({'error': str(e)})


def read_ppt_content(path: str) -> str:
    """PowerPoint(.pptx) 파일의 슬라이드 텍스트와 발표자 노트를 읽어 반환합니다."""
    try:
        ns = {'a': _A}
        with zipfile.ZipFile(path) as z:
            namelist = z.namelist()
            slide_files = sorted(
                [f for f in namelist if f.startswith('ppt/slides/slide') and f.endswith('.xml')]
            )
            slides = []
            for i, sf in enumerate(slide_files):
                root = ET.parse(z.open(sf)).getroot()
                texts = [t.text.strip() for t in root.findall('.//a:t', ns)
                         if t.text and t.text.strip()]
                num = os.path.basename(sf).replace('slide', '').replace('.xml', '')
                nf = f'ppt/notesSlides/notesSlide{num}.xml'
                notes_text = ''
                if nf in namelist:
                    nr = ET.parse(z.open(nf)).getroot()
                    notes_text = ' '.join(
                        t.text.strip() for t in nr.findall('.//a:t', ns)
                        if t.text and t.text.strip()
                    )
                slides.append({'slide': i + 1, 'content': texts, 'notes': notes_text})
            return json.dumps({'slides': slides, 'total_slides': len(slides)}, ensure_ascii=False)
    except Exception as e:
        return json.dumps({'error': str(e)})


# ── 텍스트 파일 ───────────────────────────────────────────────

def read_file(path: str, encoding: str = "utf-8") -> str:
    """텍스트 파일을 읽어 내용을 반환합니다.
    로그 파일, 설정 파일, CSV 등에 사용합니다."""
    try:
        content = Path(path).read_text(encoding=encoding, errors="replace")
        lines = content.splitlines()
        if len(lines) > 500:
            return "\n".join(lines[:500]) + f"\n... (이하 {len(lines)-500}줄 생략)"
        return content or "(빈 파일)"
    except Exception as e:
        return f"파일 읽기 실패: {e}"


def write_file(path: str, content: str, append: bool = False,
               encoding: str = "utf-8") -> str:
    """텍스트 파일에 내용을 씁니다.
    append=True 시 기존 내용 뒤에 추가합니다.
    시스템 핵심/자동실행 경로는 차단되고, 기존 파일 덮어쓰기 시 자동 백업합니다."""
    from agent.tools._safety import is_protected_path, backup_file
    if is_protected_path(path):
        return json.dumps({"blocked": True,
                           "error": f"시스템 보호 경로에는 쓸 수 없습니다: {path}"},
                          ensure_ascii=False)
    try:
        p = Path(path)
        p.parent.mkdir(parents=True, exist_ok=True)
        # 덮어쓰기(append 아님)로 기존 파일을 지우기 전 백업
        backup = backup_file(path) if (not append and p.exists()) else None
        mode = "a" if append else "w"
        with p.open(mode, encoding=encoding) as f:
            f.write(content)
        return json.dumps({"path": path, "size_bytes": p.stat().st_size,
                           "mode": "append" if append else "write",
                           "backup": backup}, ensure_ascii=False)
    except Exception as e:
        return json.dumps({"error": str(e)})


_OFFICE_EXTS = {".docx", ".xlsx", ".pptx", ".doc", ".xls", ".ppt", ".pdf"}


def office_locate_file(name: str, max_results: int = 20) -> str:
    """이름(부분 일치)으로 로컬 Office 문서를 찾습니다(OneDrive·SharePoint 동기화 폴더 포함).
    클라우드 문서를 브라우저로 편집하기 전에, 동기화된 로컬 사본을 찾아 COM으로 편집하기 위한 라운드트립 용도.
    OneDrive/SharePoint 동기화 루트와 바탕화면·문서·다운로드를 탐색합니다."""
    roots: list[str] = []
    for env in ("OneDrive", "OneDriveCommercial", "OneDriveConsumer"):
        v = os.environ.get(env)
        if v:
            roots.append(v)
    home = os.path.expanduser("~")
    for sub in ("Desktop", "Documents", "Downloads", "바탕 화면", "문서"):
        roots.append(os.path.join(home, sub))

    # 중복 루트 제거(정규화)
    seen_roots, uniq_roots = set(), []
    for r in roots:
        try:
            rr = os.path.normcase(os.path.abspath(r))
        except Exception:
            continue
        if rr in seen_roots or not os.path.isdir(r):
            continue
        seen_roots.add(rr)
        uniq_roots.append(r)

    needle = name.lower()
    matches, seen_files = [], set()
    for root in uniq_roots:
        for dirpath, dirs, files in os.walk(root):
            depth = dirpath[len(root):].count(os.sep)
            if depth > 4:
                dirs[:] = []
                continue
            dirs[:] = [d for d in dirs
                       if not d.startswith(".") and d.lower() not in ("node_modules", ".git", "appdata")]
            for f in files:
                if os.path.splitext(f)[1].lower() in _OFFICE_EXTS and needle in f.lower():
                    full = os.path.join(dirpath, f)
                    nc = os.path.normcase(full)
                    if nc in seen_files:
                        continue
                    seen_files.add(nc)
                    try:
                        mt = os.path.getmtime(full)
                    except Exception:
                        mt = 0
                    matches.append((mt, full))
            if len(matches) >= max_results * 5:
                break

    matches.sort(reverse=True)
    result = [{"path": p, "modified": mt} for mt, p in matches[:max_results]]
    return json.dumps({"query": name, "count": len(result), "matches": result,
                       "hint": ("로컬 사본이 있으면 word_edit_text/excel_set_cells(COM)로 편집하세요"
                                "(브라우저 편집보다 정확). 없으면 office_web_open을 사용하세요.")},
                      ensure_ascii=False)


MANIFEST = [
    {
        "name": "get_excel_sheet_names",
        "label": "Excel 시트 목록",
        "schema": {
            "type": "function",
            "function": {
                "name": "get_excel_sheet_names",
                "description": "Excel 파일의 시트 이름 목록을 반환합니다.",
                "parameters": {
                    "type": "object",
                    "properties": {"path": {"type": "string"}},
                    "required": ["path"]
                }
            }
        },
        "handler": lambda a: get_excel_sheet_names(a["path"])
    },
    {
        "name": "read_excel",
        "label": "Excel 읽기",
        "schema": {
            "type": "function",
            "function": {
                "name": "read_excel",
                "description": "Excel 파일을 읽어 JSON 데이터로 반환합니다. 첫 행이 헤더입니다.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "path": {"type": "string"},
                        "sheet": {"description": "시트 번호(0부터) 또는 시트 이름"},
                        "max_rows": {"type": "integer", "description": "최대 행 수 (기본 200)"}
                    },
                    "required": ["path"]
                }
            }
        },
        "handler": lambda a: read_excel(a["path"], a.get("sheet", 0), a.get("max_rows", 200))
    },
    {
        "name": "write_excel",
        "label": "Excel 쓰기",
        "schema": {
            "type": "function",
            "function": {
                "name": "write_excel",
                "description": "딕셔너리 리스트를 Excel 파일로 저장합니다.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "path": {"type": "string"},
                        "data": {"type": "array", "items": {"type": "object"}},
                        "sheet_name": {"type": "string"}
                    },
                    "required": ["path", "data"]
                }
            }
        },
        "handler": lambda a: write_excel(a["path"], a["data"], a.get("sheet_name", "Sheet1"))
    },
    {
        "name": "append_excel_row",
        "label": "Excel 행 추가",
        "schema": {
            "type": "function",
            "function": {
                "name": "append_excel_row",
                "description": "Excel 파일의 마지막에 행을 추가합니다. 파일이 없으면 새로 생성합니다.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "path": {"type": "string"},
                        "row_data": {"type": "object", "description": "헤더: 값 딕셔너리"}
                    },
                    "required": ["path", "row_data"]
                }
            }
        },
        "handler": lambda a: append_excel_row(a["path"], a["row_data"], a.get("sheet", 0))
    },
    {
        "name": "read_word",
        "label": "Word 읽기",
        "schema": {
            "type": "function",
            "function": {
                "name": "read_word",
                "description": "Word(.docx) 파일의 텍스트를 추출합니다.",
                "parameters": {
                    "type": "object",
                    "properties": {"path": {"type": "string"}},
                    "required": ["path"]
                }
            }
        },
        "handler": lambda a: read_word(a["path"])
    },
    {
        "name": "append_word",
        "label": "Word 내용 추가",
        "schema": {
            "type": "function",
            "function": {
                "name": "append_word",
                "description": "Word 파일에 내용을 추가합니다. 파일이 없으면 새로 생성합니다.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "path": {"type": "string"},
                        "text": {"type": "string"},
                        "heading": {"type": "string", "description": "섹션 제목 (선택)"}
                    },
                    "required": ["path", "text"]
                }
            }
        },
        "handler": lambda a: append_word(a["path"], a["text"], a.get("heading", ""))
    },
    {
        "name": "write_word",
        "label": "Word 문서 작성",
        "schema": {
            "type": "function",
            "function": {
                "name": "write_word",
                "description": (
                    "마크다운 텍스트를 서식이 살아있는 진짜 Word(.docx) 문서로 저장합니다. "
                    "제목(#)·목록(-, 1.)·굵게(**)·표(|)를 Word 서식으로 변환합니다. "
                    "★ .docx 파일을 만들 때는 반드시 이 도구를 사용하세요. "
                    "write_file로 .docx 경로에 텍스트를 쓰면 열 수 없는 깨진 파일이 됩니다."
                ),
                "parameters": {
                    "type": "object",
                    "properties": {
                        "path": {"type": "string", "description": ".docx 파일 경로"},
                        "content": {"type": "string", "description": "마크다운 형식의 본문"},
                        "title": {"type": "string", "description": "문서 제목 (선택)"}
                    },
                    "required": ["path", "content"]
                }
            }
        },
        "handler": lambda a: write_word(a["path"], a["content"], a.get("title", ""))
    },
    {
        "name": "read_pdf",
        "label": "PDF 읽기",
        "schema": {
            "type": "function",
            "function": {
                "name": "read_pdf",
                "description": "PDF 파일에서 텍스트를 추출합니다. pages 예: '1', '1-3', '2,4'",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "path": {"type": "string"},
                        "pages": {"type": "string", "description": "페이지 범위 (생략 시 전체)"}
                    },
                    "required": ["path"]
                }
            }
        },
        "handler": lambda a: read_pdf(a["path"], a.get("pages", ""))
    },
    {
        "name": "read_file",
        "label": "파일 읽기",
        "schema": {
            "type": "function",
            "function": {
                "name": "read_file",
                "description": "텍스트 파일을 읽어 내용을 반환합니다. 로그, 설정 파일, CSV 등에 사용합니다.",
                "parameters": {
                    "type": "object",
                    "properties": {"path": {"type": "string"}},
                    "required": ["path"]
                }
            }
        },
        "handler": lambda a: read_file(a["path"])
    },
    {
        "name": "write_file",
        "label": "파일 쓰기",
        "schema": {
            "type": "function",
            "function": {
                "name": "write_file",
                "description": "텍스트 파일(.txt, .md, .csv, .log 등)에 내용을 씁니다. append=true 시 기존 내용 뒤에 추가합니다. ★ .docx는 write_word, .xlsx는 write_excel을 사용하세요 — 이 도구로 쓰면 깨집니다.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "path": {"type": "string"},
                        "content": {"type": "string"},
                        "append": {"type": "boolean"}
                    },
                    "required": ["path", "content"]
                }
            }
        },
        "handler": lambda a: write_file(a["path"], a["content"], a.get("append", False))
    },
    {
        "name": "office_locate_file",
        "label": "로컬 Office 파일 찾기",
        "schema": {
            "type": "function",
            "function": {
                "name": "office_locate_file",
                "description": (
                    "이름(부분 일치)으로 로컬 Office 문서를 찾습니다(OneDrive/SharePoint 동기화 폴더 포함). "
                    "클라우드 문서를 브라우저로 편집하기 전에, 동기화된 로컬 사본을 찾아 COM 도구로 편집하는 "
                    "라운드트립에 사용하세요(브라우저 편집보다 정확)."
                ),
                "parameters": {
                    "type": "object",
                    "properties": {
                        "name": {"type": "string", "description": "파일명 일부(확장자 없이도 가능)"},
                        "max_results": {"type": "integer", "description": "최대 결과 수(기본 20)"},
                    },
                    "required": ["name"],
                },
            },
        },
        "handler": lambda a: office_locate_file(a["name"], a.get("max_results", 20)),
    },
    {
        "name": "read_word_comments",
        "label": "Word 검토 메모",
        "schema": {
            "type": "function",
            "function": {
                "name": "read_word_comments",
                "description": "Word(.docx) 파일의 검토 메모(Comments)를 읽어 반환합니다. 작성자·날짜·내용 포함.",
                "parameters": {
                    "type": "object",
                    "properties": {"path": {"type": "string", "description": ".docx 파일 경로"}},
                    "required": ["path"]
                }
            }
        },
        "handler": lambda a: read_word_comments(a["path"])
    },
    {
        "name": "read_word_track_changes",
        "label": "Word 수정 추적",
        "schema": {
            "type": "function",
            "function": {
                "name": "read_word_track_changes",
                "description": "Word(.docx) 파일의 수정 추적(Track Changes) 삽입·삭제 내역을 반환합니다.",
                "parameters": {
                    "type": "object",
                    "properties": {"path": {"type": "string", "description": ".docx 파일 경로"}},
                    "required": ["path"]
                }
            }
        },
        "handler": lambda a: read_word_track_changes(a["path"])
    },
    {
        "name": "read_excel_comments",
        "label": "Excel 셀 메모",
        "schema": {
            "type": "function",
            "function": {
                "name": "read_excel_comments",
                "description": "Excel(.xlsx) 파일의 셀 메모(Comments)를 읽어 반환합니다.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "path": {"type": "string", "description": ".xlsx 파일 경로"},
                        "sheet": {"type": "string", "description": "시트 이름 (생략 시 전체)"}
                    },
                    "required": ["path"]
                }
            }
        },
        "handler": lambda a: read_excel_comments(a["path"], a.get("sheet", ""))
    },
    {
        "name": "read_ppt_content",
        "label": "PPT 슬라이드 읽기",
        "schema": {
            "type": "function",
            "function": {
                "name": "read_ppt_content",
                "description": "PowerPoint(.pptx) 파일의 슬라이드 텍스트와 발표자 노트를 읽어 반환합니다.",
                "parameters": {
                    "type": "object",
                    "properties": {"path": {"type": "string", "description": ".pptx 파일 경로"}},
                    "required": ["path"]
                }
            }
        },
        "handler": lambda a: read_ppt_content(a["path"])
    },
]
